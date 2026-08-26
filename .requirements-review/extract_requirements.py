from pathlib import Path
from docx import Document

source = Path(r"E:\计算机专业考研择校平台需求规格说明书.docx")
out = Path(r"C:\Users\Lenovo\grad-school-platform\.requirements-review\requirements.txt")
document = Document(source)
lines = []
for paragraph in document.paragraphs:
    text = paragraph.text.strip()
    if text:
        lines.append(text)
for table_index, table in enumerate(document.tables, start=1):
    lines.append(f"[TABLE {table_index}]")
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append(" | ".join(cells))
out.write_text("\n".join(lines), encoding="utf-8")
print(out)
